from random import choice
import unittest
import docx
import os

# Lorem Ipsum based pre-defined test cases
TITLES = [
    "Amet mollit voluptate velit officia sit ut eu eiusmod adipisicing fugiat consequat aute occaecat. Aute proident velit proident amet. Cupidatat velit non nulla nulla occaecat proident nulla mollit velit duis Lorem non.",
    "Ipsum occaecat laborum aute aute veniam aliquip labore aliqua commodo enim aute in sit sint. Qui do officia officia eu elit laboris nulla aliquip. Exercitation do fugiat do nulla esse ullamco irure ullamco ea. Commodo consectetur nulla eiusmod proident labore amet veniam aliqua labore. Officia culpa laborum ullamco deserunt ad esse minim ad in id non ipsum. Excepteur reprehenderit quis ea excepteur consequat et nulla nulla ad id. Do aliquip et nulla consequat laboris eu fugiat.",
    "Cupidatat proident fugiat voluptate ipsum veniam. Commodo sunt nisi elit magna ullamco. Non et ea non ad officia anim pariatur. Laboris nostrud deserunt et et deserunt duis anim ex elit fugiat qui ex voluptate. Exercitation eu in est adipisicing cupidatat consequat irure nulla ad labore. Minim esse exercitation Lorem laboris.",
    "Aute anim labore ullamco ipsum. Aliqua eu Lorem consequat cupidatat sunt quis. Ex in pariatur eiusmod do pariatur cillum laborum velit velit velit ea exercitation culpa nostrud. Laboris ea aliquip ullamco qui aliqua ipsum aute aute aliqua tempor irure. Irure nulla aliqua est ad deserunt labore qui adipisicing anim sit nostrud. Tempor ea est quis eu irure qui Lorem elit magna laborum.",
    "Laborum proident nisi amet deserunt ad aute eiusmod. Do est aliquip commodo excepteur irure qui nostrud minim voluptate esse Lorem irure. Consectetur do minim ex minim esse ut.",
    "Ipsum cupidatat non ad dolore irure esse ullamco. Dolore elit aliqua cupidatat velit sit quis minim. Nisi magna culpa id eu esse fugiat nulla et pariatur nisi anim officia consectetur dolore. Qui Lorem ea anim magna enim Lorem pariatur in laborum labore nisi laborum sint sunt. Officia enim laboris esse ad laborum laborum quis deserunt exercitation elit. Labore aute qui deserunt reprehenderit ipsum do do voluptate excepteur cillum.",
    "Amet exercitation ullamco non sit velit laboris officia dolor ad officia ex elit commodo. Officia ex sint quis officia quis velit. Mollit reprehenderit pariatur labore commodo reprehenderit consectetur non in in veniam deserunt ad mollit.",
    "Enim sit veniam ut ut. Sunt eu commodo labore Lorem consectetur culpa reprehenderit non laborum amet velit deserunt. Amet ea minim quis consectetur exercitation non consequat. Commodo ipsum laborum magna occaecat ipsum consequat. Nulla eu laborum ea et veniam et sint consectetur culpa consequat.",
    "Aliquip esse velit dolore pariatur deserunt sunt voluptate ipsum. Dolore Lorem sint incididunt proident quis qui commodo voluptate. Et culpa laborum exercitation sit eiusmod adipisicing irure mollit nostrud anim cupidatat. Sint proident ea consequat id incididunt ad minim velit ex cillum.",
    "Consectetur sunt magna aliquip laborum in cupidatat et pariatur dolore anim quis est. Excepteur labore irure proident aliquip Lorem adipisicing ut pariatur. Laborum nulla Lorem sunt mollit. Quis anim voluptate incididunt culpa.",
    "Cupidatat in occaecat labore quis reprehenderit anim esse dolor esse. Mollit officia fugiat officia esse ipsum est irure dolore tempor dolor adipisicing incididunt. Ut deserunt dolore consectetur velit ad. Nulla est cupidatat nulla anim sint ipsum consequat incididunt est aliqua elit. Eu laborum ipsum non labore minim pariatur elit nisi. Adipisicing ad do laboris incididunt cupidatat duis occaecat cillum cillum eiusmod ea deserunt ex.",
    "Anim proident amet laborum aliquip in velit esse aliquip in dolore eu irure aliqua. Elit qui commodo dolore pariatur exercitation aute ea ex aliquip tempor minim officia. Ut commodo fugiat dolor officia laborum pariatur incididunt anim aliqua. Eiusmod minim ut quis nostrud in. Mollit eu consectetur id ea excepteur commodo ut ex deserunt dolor laborum ipsum aliqua eu. Do consequat amet consequat eu quis ullamco enim est ad occaecat. Tempor dolore do Lorem reprehenderit fugiat fugiat consequat.",
    "Minim aliquip eiusmod qui ad veniam quis magna velit. Amet qui laboris sit labore. Veniam duis cillum eu irure reprehenderit voluptate ex. Mollit ut incididunt commodo et reprehenderit proident sit reprehenderit esse."
]
POSSIBLE_ANSWERS = [
    "Reprehenderit et laborum veniam cillum veniam est id et ut non id dolor proident do.",
    "Proident commodo duis ex ea aliquip sunt.",
    "Irure cillum deserunt labore ad ipsum pariatur in dolor sint consequat dolore excepteur enim.",
    "Mollit elit incididunt nulla culpa qui cillum.",
    "Consectetur minim ex consectetur est sit fugiat qui culpa duis minim.",
    "Nostrud laborum irure nisi commodo deserunt magna aliquip aliquip proident.",
    "Ea aute ullamco ex esse ipsum sunt aute nisi tempor est consectetur nostrud.",
    "Et est eu quis consequat fugiat ut eiusmod tempor culpa.",
    "Lorem Lorem nisi cillum non non duis consequat exercitation magna.",
    "Est laborum dolor aute laboris minim commodo.",
    "Irure est aute occaecat tempor.",
    "Minim officia excepteur laboris mollit ex aliquip ipsum consectetur ipsum Lorem.",
    "Quis dolor reprehenderit laborum consequat amet incididunt aliquip culpa non occaecat nisi.",
    "Non minim ad Lorem quis enim nisi laborum mollit sit."
]


import assets.file_manage as fm

class FileLoading(unittest.TestCase):
    # Setup './test.docx' file
    def setUp(self) -> None:
        self.expected_answers = {}
        self.answer_string = ""
        # Setting up the document
        testDoc = docx.Document()
        for i in range(3):
            TITLE = choice(TITLES)
            OPTIONS = [f" {choice(POSSIBLE_ANSWERS)}", f" {choice(POSSIBLE_ANSWERS)}", f" {choice(POSSIBLE_ANSWERS)}", f" {choice(POSSIBLE_ANSWERS)}"]
            ANSWER = choice([0, 1, 2, 3])
            self.expected_answers[i] = {"title": f"Câu {i+1}: {TITLE}", "answer": OPTIONS[ANSWER], "options": OPTIONS}
            if ANSWER == 0:
                self.answer_string += "A"
            elif ANSWER == 1:
                self.answer_string += "B"
            elif ANSWER == 2:
                self.answer_string += "C"
            elif ANSWER == 3:
                self.answer_string += "D"

            testDoc.add_paragraph(f'Câu {i+1}: {TITLE}')
            testDoc.add_paragraph(f'A. {OPTIONS[0]}')
            testDoc.add_paragraph(f'B. {OPTIONS[1]}')
            testDoc.add_paragraph(f'C. {OPTIONS[2]}')
            testDoc.add_paragraph(f'D. {OPTIONS[3]}')
        testDoc.save('test.docx')

    def test_read_and_format(self):
        try:
            txt = fm.convertDoc2Txt(r"./test.docx")
            final = fm.formatter(txt)
        except Exception as e:
            self.fail(f"Exception Raised: {e}")
        
    def test_filtering(self):
        try:

            txt = fm.convertDoc2Txt(r"./test.docx")
            final = fm.formatter(txt)
        except Exception as e:
            self.fail(f"Exception Raised: {e}")

        self.assertDictEqual(fm.filterer(self.answer_string, "2", final), self.expected_answers)

    def tearDown(self) -> None:
        # Remove the test file
        os.remove(r'./test.docx')

